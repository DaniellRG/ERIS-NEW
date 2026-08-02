# -*- coding: utf-8 -*-
"""document_generator.py — Genera documentos Word profesionales con python-docx.
Portada, secciones, listas, citas, formato. Crea la carpeta automaticamente."""
import json
import os
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

_DATA_DIR = Path(__file__).resolve().parent.parent / "data"
_CREATED_FILE = _DATA_DIR / "documents_created.json"
_WORKING_FILE = _DATA_DIR / "working_document.json"


def document_generator(parameters: dict, player=None) -> str:
    action = parameters.get("action", "create").lower()
    if action == "create":
        return _create_document(parameters, player)
    elif action == "check_content":
        return _check_content(parameters)
    elif action == "working_doc":
        return _get_working_doc(parameters)
    elif action == "list_templates":
        return _list_templates()
    elif action in ("to_pdf", "convert"):
        return _convert_to_pdf(parameters)
    elif action == "to_epub":
        return _convert_to_epub(parameters)
    return "Acciones: create (crear doc), check_content (leer doc), working_doc (doc en curso), to_pdf (convertir a PDF), list_templates (plantillas)."


def _create_document(parameters: dict, player=None) -> str:
    title = parameters.get("title", "Documento")
    subtitle = parameters.get("subtitle", "")
    author = parameters.get("author", "ERIS")
    sections_raw = parameters.get("sections", "[]")
    full_text = parameters.get("content", "") or parameters.get("full_text", "")
    filename = parameters.get("filename", "")
    output_path_raw = parameters.get("output_path", "")
    convert_pdf = str(parameters.get("convert_pdf", "false")).lower() == "true"

    if not _DOCX_OK:
        return "Error: Falta python-docx. Instala con: pip install python-docx"

    try:
        doc = Document()
        _set_styles(doc)
        _add_cover_page(doc, title, subtitle, author)

        if full_text:
            _add_full_text(doc, full_text)
        else:
            sections = json.loads(sections_raw) if isinstance(sections_raw, str) else sections_raw
            for sec in sections:
                heading = sec.get("heading", "")
                content = sec.get("content", "")
                style = sec.get("style", "normal")
                items = sec.get("items", [])
                level = sec.get("level", 1)

                if heading:
                    doc.add_heading(heading, level=min(level, 3))

                if items:
                    for item in items:
                        if style == "numbered":
                            doc.add_paragraph(item, style='List Number')
                        else:
                            doc.add_paragraph(item, style='List Bullet')
                elif content:
                    for para_text in content.split("\n"):
                        para_text = para_text.strip()
                        if not para_text:
                            continue
                        if style == "quote":
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            run = p.add_run(para_text)
                            run.italic = True
                            run.font.size = Pt(10)
                            p.paragraph_format.left_indent = Cm(1)
                        elif style == "center":
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.add_run(para_text)
                        else:
                            p = doc.add_paragraph(para_text)
                            p.paragraph_format.space_after = Pt(6)

        output_dir = _resolve_output_dir(output_path_raw)
        output_dir.mkdir(parents=True, exist_ok=True)

        safe_name = "".join(c for c in title if c.isalnum() or c in " _-").strip().replace(" ", "_")
        if not filename:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{safe_name}_{ts}.docx"
        if not filename.endswith(".docx"):
            filename += ".docx"

        docx_path = output_dir / filename
        doc.save(str(docx_path))

        _track_created("Word", docx_path, title)
        _save_working_doc(title, str(docx_path))

        result = f"Documento creado: {docx_path} ({os.path.getsize(docx_path)} bytes)"

        if convert_pdf:
            pdf_result = _do_convert_to_pdf(docx_path, output_dir)
            result += f"\n{pdf_result}"

        if player:
            player.write_log(result)

        return result

    except json.JSONDecodeError:
        return "Error: sections debe ser un JSON valido."
    except Exception as e:
        return f"Error creando documento: {e}"


def _track_created(fmt: str, path: Path, title: str):
    try:
        created = []
        if _CREATED_FILE.exists():
            created = json.loads(_CREATED_FILE.read_text(encoding="utf-8"))
        created.append({
            "format": fmt, "path": str(path),
            "title": title[:80], "time": datetime.now().isoformat(),
        })
        if len(created) > 100:
            created = created[-100:]
        _CREATED_FILE.parent.mkdir(parents=True, exist_ok=True)
        _CREATED_FILE.write_text(json.dumps(created, indent=2, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _save_working_doc(title: str, path: str):
    try:
        _WORKING_FILE.parent.mkdir(parents=True, exist_ok=True)
        _WORKING_FILE.write_text(json.dumps({
            "title": title, "path": path, "time": datetime.now().isoformat(),
        }, indent=2, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def _get_working_doc(parameters: dict) -> str:
    if not _WORKING_FILE.exists():
        return "No hay documento en curso."
    try:
        data = json.loads(_WORKING_FILE.read_text(encoding="utf-8"))
        return "Documento en curso:\n  Titulo: {}\n  Ruta: {}\n  Ultima modificacion: {}".format(
            data.get("title", "?"), data.get("path", "?"), data.get("time", "?")[:16])
    except Exception:
        return "Error leyendo documento en curso."


def _check_content(parameters: dict) -> str:
    path_str = parameters.get("path", "")
    if not path_str:
        if _WORKING_FILE.exists():
            try:
                data = json.loads(_WORKING_FILE.read_text(encoding="utf-8"))
                path_str = data.get("path", "")
            except Exception:
                pass
        if not path_str:
            created = _load_created_list()
            if created:
                path_str = created[-1].get("path", "")
    if not path_str or not Path(path_str).exists():
        return "No se encontro el documento. Especifica path o crea un documento primero."
    path = Path(path_str)
    try:
        if path.suffix.lower() == ".docx":
            from docx import Document
            doc = Document(str(path))
            paras = [p.text for p in doc.paragraphs]
        elif path.suffix.lower() == ".txt":
            paras = path.read_text(encoding="utf-8").split("\n")
        elif path.suffix.lower() == ".md":
            paras = path.read_text(encoding="utf-8").split("\n")
        else:
            return "Formato no soportado para lectura: {}".format(path.suffix)
        lines = [p for p in paras if p.strip()]
        summary = "Documento: {} ({} parrafos con contenido, {} KB)".format(
            path.name, len(lines), round(path.stat().st_size / 1024, 1))
        preview = "\n".join(lines[:30])
        if len(lines) > 30:
            preview += "\n... ({} parrafos mas)".format(len(lines) - 30)
        return summary + "\n\n" + preview
    except Exception as e:
        return "Error leyendo documento: {}".format(str(e)[:80])


def _load_created_list() -> list:
    if _CREATED_FILE.exists():
        try:
            return json.loads(_CREATED_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return []


def _add_full_text(doc, text: str):
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("1. ") or line.startswith("1) "):
            p = doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('"') and line.endswith('"'):
            p = doc.add_paragraph()
            run = p.add_run(line[1:-1])
            run.italic = True
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Cm(1)
        else:
            para = line
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line or next_line.startswith("#") or next_line.startswith("- "):
                    break
                para += " " + next_line
                i += 1
            doc.add_paragraph(para)


def _set_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    for level in range(1, 5):
        try:
            h = doc.styles[f'Heading {level}']
            h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
            if level == 1:
                h.font.size = Pt(18)
            elif level == 2:
                h.font.size = Pt(14)
            else:
                h.font.size = Pt(12)
        except Exception:
            pass


def _add_cover_page(doc, title, subtitle, author):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    if subtitle:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Por: {author}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%d de %B de %Y"))
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def _resolve_output_dir(path_raw: str) -> Path:
    if path_raw:
        p = Path(path_raw)
        if not p.is_absolute():
            try:
                from actions.path_helper import get_desktop_path
                desk = Path(get_desktop_path())
            except Exception:
                desk = Path.home() / "Desktop"
            p = desk / path_raw
        return p
    try:
        from actions.path_helper import get_desktop_path
        desk = Path(get_desktop_path())
    except Exception:
        desk = Path.home() / "Desktop"
    return desk / "ERIS_Documentos"


def _convert_to_pdf(parameters: dict) -> str:
    path = parameters.get("path", "")
    output = parameters.get("output", "")
    if not path:
        return "Error: Requiero 'path' al archivo .docx"
    src = Path(path)
    if not src.exists():
        return f"Error: No existe el archivo: {src}"
    if src.suffix.lower() not in (".docx", ".doc"):
        return f"Error: Solo soporto .docx, no {src.suffix}"
    output_dir = Path(output) if output else src.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = output_dir / (src.stem + ".pdf")
    return _do_convert_to_pdf(src, output_dir, pdf_path.name)


def _do_convert_to_pdf(docx_path: Path, output_dir: Path, filename: str = "") -> str:
    if not filename:
        filename = docx_path.stem + ".pdf"
    pdf_path = output_dir / filename
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        try:
            wdFormatPDF = 17
            doc = word.Documents.Open(str(docx_path.resolve()))
            doc.SaveAs(str(pdf_path.resolve()), FileFormat=wdFormatPDF)
            doc.Close()
            return f"PDF creado: {pdf_path} ({os.path.getsize(pdf_path)} bytes)"
        finally:
            word.Quit()
    except Exception as e:
        return f"Error convirtiendo a PDF con Word: {e}"


def _list_templates() -> str:
    return (
        "Acciones de document_generator:\n\n"
        "1. create: Crear documento Word\n"
        "   title (str): Titulo del documento\n"
        "   subtitle (str): Subtitulo opcional\n"
        "   author (str): Autor (default: ERIS)\n"
        "   sections (str): JSON con secciones estructuradas\n"
        "     [{\"heading\":\"Titulo\",\"content\":\"Texto\",\"style\":\"normal|bullet|numbered|quote|center\",\"items\":[],\"level\":1}]\n"
        "   full_text (str): Texto plano largo con marcado simple:\n"
        "     # Titulo nivel 1\n"
        "     ## Titulo nivel 2\n"
        "     ### Titulo nivel 3\n"
        "     - item de lista\n"
        "     Parrafos separados por linea en blanco\n"
        "     \"cita entre comillas\"\n"
        "   convert_pdf (str): \"true\" para crear tambien PDF\n"
        "   filename (str): Nombre del archivo\n"
        "   output_path (str): Carpeta destino (default: Desktop\\ERIS_Documentos)\n\n"
        "2. to_pdf: Convertir DOCX a PDF\n"
        "   path (str): Ruta al archivo .docx\n"
        "   output (str): Carpeta destino (default: mismo lugar)\n\n"
        "3. check_content: Inspeccionar contenido de un documento\n"
        "   path (str): Ruta al archivo (opcional, usa el ultimo si no se especifica)\n\n"
        "4. working_doc: Ver el documento en curso\n\n"
        "5. list_templates: Esta ayuda\n\n"
        "Ejemplo full_text:\n"
        '  document_generator action=create title="Mi Libro" full_text="# Capitulo 1\\n\\nTexto largo aqui..."'
    )
